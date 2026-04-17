from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_title(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    if level == 0:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'D9E2F3')

    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)

    return table

doc = Document()

title = add_title(doc, '三国纵横·消战天下 - 需求文档', 0)
title.runs[0].font.size = Pt(24)
title.runs[0].font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph('版本：V1.0').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('日期：2024年').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_title(doc, '1. 应用概述', 1)
doc.add_paragraph('• 应用名称：三国纵横·消战天下')
doc.add_paragraph('• 应用描述：一款融合三消玩法与卡牌将领养成的策略手游。玩家扮演乱世主公，通过三消战斗参与三国历史战役，消除棋子产出资源，养成将领，推进主线剧情最终一统天下。')

add_title(doc, '2. 用户与使用场景', 1)

add_title(doc, '2.1 目标用户', 2)
doc.add_paragraph('• 年龄18-35岁，喜欢策略游戏、三国题材及消除类休闲游戏')
doc.add_paragraph('• 游戏时间碎片化，单局5-15分钟')
doc.add_paragraph('• 有收集养成倾向，喜欢阵容搭配和资源管理')

add_title(doc, '2.2 核心使用场景', 2)
doc.add_paragraph('• 挑战战役关卡（三消战斗）→ 获得资源 → 养成将领 → 挑战更强关卡（循环）')
doc.add_paragraph('• 单机或异步竞争，不强制实时PVP')

add_title(doc, '3. 页面结构与功能说明', 1)

add_title(doc, '3.1 整体页面结构', 2)
doc.add_paragraph('【页面导航流程图】')
doc.add_paragraph()

doc.add_paragraph('详细页面结构：')
doc.add_paragraph('三国纵横·消战天下')
doc.add_paragraph('├── 主菜单页')
doc.add_paragraph('│   ├── 游戏Logo与背景动画')
doc.add_paragraph('│   ├── 「开始游戏」按钮')
doc.add_paragraph('│   ├── 「设置」按钮 → 设置面板')
doc.add_paragraph('│   └── 「退出」按钮')
doc.add_paragraph('├── 战役关卡选择页')
doc.add_paragraph('├── 战斗页（三消核心）')
doc.add_paragraph('├── 关卡结算页')
doc.add_paragraph('├── 将领页')
doc.add_paragraph('├── 抽卡页')
doc.add_paragraph('└── 商店页')

add_title(doc, '3.2 主菜单页', 2)
doc.add_paragraph('核心功能：')
doc.add_paragraph('• 展示游戏Logo及动态背景动画（三国战场风格）')
doc.add_paragraph('• 提供「开始游戏」按钮，点击直接进入战役关卡选择页')
doc.add_paragraph('• 提供「设置」按钮，点击弹出设置面板')
doc.add_paragraph('• 提供「退出」按钮，点击确认后退出游戏')
doc.add_paragraph('• 提供存档的注册起名功能')

doc.add_paragraph()
doc.add_paragraph('设置面板：')
doc.add_paragraph('• 音乐开关（默认开启）')
doc.add_paragraph('• 音效开关（默认开启）')
doc.add_paragraph('• 音量调节滑块（0-100）')

add_title(doc, '3.3 底部导航栏', 2)
doc.add_paragraph('功能：在战役、将领、抽卡、商店页面之间切换')
doc.add_paragraph('导航项：')
doc.add_paragraph('• 战役：进入战役关卡选择页')
doc.add_paragraph('• 将领：进入将领页')
doc.add_paragraph('• 抽卡：进入抽卡页')
doc.add_paragraph('• 商店：进入商店页')

add_title(doc, '3.4 战役关卡选择页', 2)
doc.add_paragraph('章节列表（按三国编年史顺序）：')
doc.add_paragraph('• 序章：黄巾之乱（1-10关）')
doc.add_paragraph('• 第一章：诸侯讨董（11-20关）[敬请期待]')
doc.add_paragraph('• 第二章：官渡之战（21-35关）[敬请期待]')
doc.add_paragraph('• 第三章：赤壁之战（36-50关）[敬请期待]')
doc.add_paragraph('• 第四章：三分天下（51-70关）[敬请期待]')
doc.add_paragraph('• 终章：一统天下（71-80关）[敬请期待]')

add_title(doc, '3.5 战斗页（三消核心）', 2)

add_title(doc, '3.5.1 顶部HUD', 3)
doc.add_paragraph('• 主公头像、等级、主公名称')
doc.add_paragraph('• 当前生命值（数值 + 进度条）')
doc.add_paragraph('• 当前能量值（数值 + 进度条，上限100）')
doc.add_paragraph('• 设置按钮（暂停/退出关卡）')

add_title(doc, '3.5.2 8×8三消棋盘', 3)
doc.add_paragraph('棋盘规格：8×8网格，竹简卷轴风格边框')
doc.add_paragraph()
doc.add_paragraph('棋子类型（6种基础色）：')
add_table(doc,
    ['颜色', '类型', '效果'],
    [
        ['红色', '攻击', '消除时对敌方造成伤害'],
        ['绿色', '生命', '消除时回复我方生命'],
        ['黄色', '士气', '消除时增加技能能量'],
        ['蓝色', '防御', '消除时获得临时护盾'],
        ['紫色', '谋略', '消除时增加连击点数'],
        ['橙色', '粮草', '消除时掉落额外资源']
    ])

doc.add_paragraph()
doc.add_paragraph('特殊棋子：')
doc.add_paragraph('• 强化棋子：4个连珠生成，消除时触发范围效果')
doc.add_paragraph('• 炸弹棋子：5个连珠生成，消除时炸毁3×3区域')
doc.add_paragraph('• 十字消除棋子：T/L型消除生成，消除时触发十字范围消除')

add_title(doc, '3.5.3 将领技能栏', 3)
doc.add_paragraph('• 每场战斗最多携带2名将领')
doc.add_paragraph('• 能量机制：初始0，每消除1个普通棋子+1，消除特殊棋子额外+2，上限100')
doc.add_paragraph('• 点击将领头像：能量满足时释放技能，消耗对应能量')
doc.add_paragraph('• 技能无CD，仅受能量限制')

add_title(doc, '3.5.4 敌方信息栏', 3)
doc.add_paragraph('• 展示敌方名称、等级、当前生命值（数值+进度条）')
doc.add_paragraph('• 展示敌方下一回合攻击预告')
doc.add_paragraph('• 敌方生命归零则关卡胜利')

add_title(doc, '3.6 将领页', 2)
doc.add_paragraph('将领列表：')
doc.add_paragraph('• 展示已拥有将领列表，含头像、名称、阵营、稀有度、当前等级/星级')
doc.add_paragraph('• 支持按阵营、稀有度、等级排序')
doc.add_paragraph()
doc.add_paragraph('将领详情页：')
doc.add_paragraph('• 基础属性：攻击、生命')
doc.add_paragraph('• 升级功能：消耗银两和将领经验书，提升攻击/生命属性')
doc.add_paragraph('• 升星功能：消耗将领碎片、银两和经验书，提升将领品质和属性，最高5星')

add_title(doc, '3.7 抽卡页', 2)
doc.add_paragraph('• 卡池选择：普通池/限定池（限时活动）')
doc.add_paragraph('• 抽卡方式：单抽（100元宝）、十连抽（900元宝，9折优惠）')
doc.add_paragraph('• 十连抽必出至少1个稀有及以上品质将领')

add_title(doc, '3.8 商店页', 2)
doc.add_paragraph('商品分类：')
doc.add_paragraph('• 体力药水：恢复体力')
doc.add_paragraph('• 银两袋：增加银两')
doc.add_paragraph('• 木材包：增加木材')
doc.add_paragraph('• 将领经验书：用于将领升级')
doc.add_paragraph('• 万能碎片：可用于任何将领升星')

add_title(doc, '3.9 背包页', 2)
doc.add_paragraph('• 道具列表：展示当前持有道具列表（类型、数量）')
doc.add_paragraph('• 道具操作：点击道具可查看详细说明')
doc.add_paragraph('• 部分道具可直接使用（如体力药水）')

add_title(doc, '4. 业务规则与逻辑', 1)

add_title(doc, '4.1 阵营基础效果（主公被动）', 2)
add_table(doc,
    ['阵营', '主公名称', '基础效果'],
    [
        ['魏', '曹操', '每消除10个蓝色棋子，获得1层护盾（每层减伤5%，上限4层）'],
        ['蜀', '刘备', '每消除10个绿色棋子，回复5%生命'],
        ['吴', '孙权', '每消除10个红色棋子，对随机敌人造成50%攻击力的火焰伤害'],
        ['群雄', '董卓', '每消除10个紫色棋子，下次技能伤害+30%'],
        ['群雄', '张角', '每消除10个黄色棋子，获得1层「黄天之力」（每层增加技能伤害10%，上限5层）']
    ])

add_title(doc, '4.2 主公列表（初始可选1名）', 2)
add_table(doc,
    ['主公', '阵营', '初始攻击', '初始生命', '被动技能', '获得方式'],
    [
        ['曹操', '魏', '80', '1200', '护盾积累', '初始可选'],
        ['刘备', '蜀', '70', '1400', '生命回复', '初始可选'],
        ['孙权', '吴', '90', '1100', '火焰伤害', '初始可选'],
        ['董卓', '群雄', '100', '1000', '技能增伤', '初始可选'],
        ['张角', '群雄', '95', '1150', '黄天之力', '碎片合成（需张角碎片×50）']
    ])

add_title(doc, '4.3 将领列表（V1.0 部分示例）', 2)
add_table(doc,
    ['将领', '阵营', '稀有度', '初始攻击', '初始生命', '技能', '技能效果'],
    [
        ['司马懿', '魏', '传说', '85', '900', '隐忍待发', '将所有蓝色棋子转化为强化棋子，并获得2层护盾'],
        ['郭嘉', '魏', '传说', '75', '850', '奇谋妙算', '随机将10个紫色棋子变为炸弹棋子，并增加10点能量'],
        ['张飞', '蜀', '传说', '90', '800', '横扫千军', '消除最下面3行所有棋子，对敌人造成额外10%伤害'],
        ['诸葛亮', '蜀', '传说', '80', '750', '八卦阵', '将棋盘上所有棋子重新排列，确保至少3组可消除'],
        ['周瑜', '吴', '传说', '85', '700', '火烧连营', '随机消除15个红色棋子，并对敌人造成持续2回合的火焰伤害'],
        ['小乔', '吴', '传说', '65', '850', '琴音绕梁', '将所有绿色棋子转化为黄色棋子，并为友军回复15%生命'],
        ['吕布', '群雄', '传说', '100', '900', '乱世枭雄', '消除棋盘上所有同色棋子，并对敌人造成相当于消除数量×5的伤害'],
        ['貂蝉', '群雄', '传说', '70', '800', '倾国倾城', '随机将8个紫色棋子变为黄色棋子，并使敌人下一回合攻击伤害降低30%'],
        ['张宝', '群雄', '史诗', '75', '850', '地公将军', '消除棋盘上所有绿色棋子，并对敌人造成相当于消除数量×3的伤害']
    ])

add_title(doc, '4.4 经济系统', 2)
add_table(doc,
    ['货币名称', '获取方式', '消耗方式', '是否可充值'],
    [
        ['银两', '三消关卡掉落、城池税收、任务奖励', '升级城池、将领升级、购买道具', '否'],
        ['元宝', '充值、成就奖励、每日任务、活动奖励', '购买体力、抽卡、加速建造、购买稀有道具', '是'],
        ['木材', '三消关卡掉落、工坊生产、任务奖励', '升级城池建筑', '否'],
        ['将领碎片', '关卡掉落、抽卡、活动、商店购买', '将领升星、合成新将领', '间接充值'],
        ['技能点', '将领升级、活动奖励', '技能升级（后续版本）', '否']
    ])

add_title(doc, '4.5 体力系统', 2)
doc.add_paragraph('• 初始体力：100')
doc.add_paragraph('• 战役关卡消耗：10体力/次')
doc.add_paragraph('• 体力恢复：5分钟恢复1点')
doc.add_paragraph('• 体力上限可通过升级主殿提高，最高200')
doc.add_paragraph('• 体力不足时提示购买体力药水或等待恢复')
doc.add_paragraph('• 每日体力购买次数限制：5次（随VIP等级提升）')

add_title(doc, '4.6 成长数值', 2)
doc.add_paragraph('• 主公升级经验 = 100 × (当前等级^1.2)')
doc.add_paragraph('• 将领升级经验 = 50 × (当前等级^1.3)')
doc.add_paragraph('• 主公每升一级：体力上限+2，解锁新功能/建筑')
doc.add_paragraph('• 将领每升一级：攻击+5，生命+50（基础值）')
doc.add_paragraph('• 将领升星成长率：每升一星，属性成长率提升10%')

add_title(doc, '4.7 敌人AI行为', 2)
doc.add_paragraph('敌人类型：普通兵种（步兵/弓兵/骑兵）、将领BOSS、守城器械')
doc.add_paragraph('• 普通兵种：每回合对玩家造成固定伤害（10-50）')
doc.add_paragraph('• 将领BOSS：每3回合释放一次全体攻击技能')
doc.add_paragraph('• 守城器械：每回合减少玩家能量或随机消除棋子')
doc.add_paragraph()
doc.add_paragraph('行为逻辑：')
doc.add_paragraph('• 敌人生命值 > 50%：普通攻击')
doc.add_paragraph('• 敌人生命值 ≤ 50%：30%概率释放技能，否则普通攻击')
doc.add_paragraph('• 敌人难度：随章节推进逐渐提升，攻击力和生命值增加')

add_title(doc, '4.8 关卡奖励规则', 2)
doc.add_paragraph('• 首次通关：银两×500 + 木材×100 + 将领经验书×1（以第1关为基准，后续关卡奖励递增）')
doc.add_paragraph('• 重复通关：银两×200 + 木材×50 + 概率掉落碎片')
doc.add_paragraph('• 精英关：将领碎片掉落概率更高（50%概率）')
doc.add_paragraph('• BOSS关：必掉特定将领碎片（2-5个）')
doc.add_paragraph('• 章节奖励：完成整章节可获得额外奖励（元宝、稀有将领碎片）')

add_title(doc, '4.9 关卡目标类型', 2)
doc.add_paragraph('• 标准目标：击败所有敌人')
doc.add_paragraph('• 限制回合：在指定回合数内完成（如20回合内）')
doc.add_paragraph('• 保护目标：保护指定目标（如粮草车）不被摧毁')
doc.add_paragraph('• 收集目标：收集指定数量的特定棋子')

add_title(doc, '4.10 音频设计', 2)
doc.add_paragraph('背景音乐：')
doc.add_paragraph('• 主菜单：大气国风，展示游戏世界观')
doc.add_paragraph('• 城池界面：悠扬古筝+笛子，宁静国风')
doc.add_paragraph('• 战斗界面：快板鼓+二胡，紧张激昂')
doc.add_paragraph('• BOSS战：交响+人声合唱，史诗感')
doc.add_paragraph('• 结算界面：轻松愉快，庆祝胜利')
doc.add_paragraph()
doc.add_paragraph('音效：')
doc.add_paragraph('• 消除棋子：竹简翻动声')
doc.add_paragraph('• 技能释放：刀剑出鞘声+语音')
doc.add_paragraph('• 升级成功：锣声')
doc.add_paragraph('• 点击UI：毛笔点击声')
doc.add_paragraph('• 资源获得：金币掉落声')

add_title(doc, '5. 异常与边界情况', 1)
add_table(doc,
    ['场景', '异常情况', '处理方式'],
    [
        ['棋子交换', '交换后无法形成≥3连珠', '交换无效，播放错误提示动画，棋子复位'],
        ['棋盘死局', '棋盘上无任何可消除组合', '自动重新随机排列棋盘，不消耗回合'],
        ['技能释放', '能量不足', '按钮置灰，点击时提示「能量不足」'],
        ['体力不足', '挑战关卡时体力不足', '弹窗提示，提供「购买体力药水」或「等待恢复」选项'],
        ['玩家生命归零', '战斗中生命值降至0', '触发失败结算，展示失败页面'],
        ['网络中断', '战斗中断网', '本地保存当前战斗状态，重连后可继续'],
        ['升级资源不足', '银两或木材不足时点击升级', '提示「资源不足」，引导前往商店或挑战关卡'],
        ['将领碎片不足', '升星时碎片不足', '提示「碎片不足」，展示碎片获取途径'],
        ['元宝不足', '购买道具或抽卡时元宝不足', '提示「元宝不足」，引导前往充值界面'],
        ['每日购买限制', '超过每日体力购买次数限制', '提示「购买次数已达上限」'],
        ['关卡未解锁', '尝试进入未解锁的关卡', '提示「关卡未解锁」，引导完成前置关卡']
    ])

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph('—— 文档结束 ——')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('docs/PRD_需求文档_v1.0.docx')
print('Word文档已生成：docs/PRD_需求文档_v1.0.docx')